import streamlit as st
import cv2
import numpy as np
import difflib
import language_tool_python
from nltk.corpus import wordnet
from nltk.tokenize import word_tokenize
import nltk
import easyocr
from datetime import datetime
import docx

# Download NLTK data
nltk.download('punkt')
nltk.download('wordnet')

# Initialize EasyOCR reader
reader = easyocr.Reader(['en'])

# Initialize LanguageTool
tool = language_tool_python.LanguageTool('en-US')

# Define the standard answers with key points
standard_answers = {
    "Q1": [
        "An operating system (OS) is system software that manages computer hardware and software resources, providing common services for computer programs.",
        "An OS is software that provides a platform for running applications and managing hardware resources.",
        "The operating system acts as an intermediary between users and the computer hardware, enabling the execution of programs."
    ],
    "Q2": [
        "Applications of an OS include managing hardware resources, providing a user interface, executing applications, and ensuring security and access control.",
        "An OS is responsible for managing resources, facilitating user interaction through the interface, and controlling security and access.",
        "The OS serves to manage computer resources, run applications, and provide security to user data and resources."
    ]
}

# Tokenize answers and keep only significant keywords for comparison
def extract_key_phrases(answer):
    words = word_tokenize(answer.lower())
    key_phrases = [word for word in words if word not in ["a", "an", "the", "and", "is", "to", "for", "in", "on", "of"]]
    return set(key_phrases)

# Function to calculate similarity score using difflib
def calculate_similarity(answer1, answer2):
    similarity = difflib.SequenceMatcher(None, answer1, answer2).ratio()
    return similarity

# Function to check grammar using LanguageTool
def check_grammar(answer):
    matches = tool.check(answer)
    corrected_answer = language_tool_python.utils.correct(answer, matches)
    grammar_score = 1 - (len(matches) / len(answer)) if len(answer) > 0 else 0
    return grammar_score, corrected_answer

# Function to identify missing parts by comparing student's answer with standard answer
def identify_missing_parts(student_answer, standard_answer):
    student_key_phrases = extract_key_phrases(student_answer)
    standard_key_phrases = extract_key_phrases(standard_answer)
    missing_phrases = standard_key_phrases - student_key_phrases
    return list(missing_phrases)

# Function to evaluate the answer and dynamically identify mistakes
def evaluate_answer(student_answer, standard_answers, marks, length_based=False):
    if student_answer.strip().upper() == "UNATTEMPTED":
        return 0, 0, 0, 0, ["Unattempted"]

    best_score = 0
    best_mistakes = []
    best_similarity_score = 0
    best_grammar_score = 0
    best_length_score = 1 if not length_based else 0

    for standard_answer in standard_answers:
        similarity_score = calculate_similarity(student_answer, standard_answer)
        grammar_score, corrected_answer = check_grammar(student_answer)
        missing_parts = identify_missing_parts(student_answer, standard_answer)
        missing_mistakes = [f"Missing part: '{part}'" for part in missing_parts]

        if length_based:
            length_score = min(len(student_answer) / len(standard_answer), 1)
        else:
            length_score = 1

        total_score = (similarity_score * 0.4 + grammar_score * 0.2 + length_score * 0.1) * marks

        if total_score > best_score:
            best_score = total_score
            best_similarity_score = similarity_score * 0.4 * marks
            best_grammar_score = grammar_score * 0.2 * marks
            best_length_score = length_score * 0.1 * marks if length_based else 0
            best_mistakes = missing_mistakes

    return best_score, best_similarity_score, best_grammar_score, best_length_score, best_mistakes

# Function to process image and extract answers separated by "END"
def process_image(image_path):
    result = reader.readtext(image_path)
    extracted_text = "\n".join([item[1] for item in result])
    answers = [answer.strip() for answer in extracted_text.split("END") if answer.strip()]
    return answers

# Function to save text into a Word file
def save_to_word(answers, filename):
    doc = docx.Document()
    for answer in answers:
        doc.add_paragraph(answer)
        doc.add_paragraph("\n--------------------------------------------------")
    doc.save(filename)

# Streamlit app
st.title("Answer Evaluation System")

uploaded_file = st.file_uploader("Upload an Image", type=["png", "jpg", "jpeg"])
if uploaded_file:
    with open(f"uploaded_image.png", "wb") as f:
        f.write(uploaded_file.getbuffer())
    st.success("Image uploaded successfully!")

    # Initialize session state for tracking submit/re-evaluate state
    if "submitted" not in st.session_state:
        st.session_state.submitted = False

    if st.button("Submit for Evaluation", disabled=st.session_state.submitted):
        answers = process_image("uploaded_image.png")

        # Save extracted text to Word file
        word_filename = f"output_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
        save_to_word(answers, word_filename)
        st.success(f"Text extracted and saved as {word_filename}.")

        # Store initial evaluation scores and mistakes
        st.session_state.submitted = True
        st.session_state.answers = answers

        total_score = 0
        all_mistakes = {"Q1": [], "Q2": []}

        # Displaying scores, mistakes, and achievements
        def display_with_box(title, content, color):
            st.markdown(
                f"""
                <div style='border: 1px solid {color}; border-radius: 10px; padding: 10px; margin: 10px 0;'>
                    <strong>{title}</strong>: {content}
                </div>
                """, unsafe_allow_html=True
            )

        if len(answers) > 0:
            q1_score, q1_similarity, q1_grammar, _, q1_mistakes = evaluate_answer(answers[0], standard_answers['Q1'], 2)
            total_score += q1_score
            all_mistakes["Q1"].extend(q1_mistakes)
            
            display_with_box("Q1 Submitted Score", f"{q1_score:.2f}", "green")
            display_with_box("Q1 Submitted Similarity Marks", f"{q1_similarity:.2f}", "blue")
            display_with_box("Q1 Submitted Grammar Marks", f"{q1_grammar:.2f}", "blue")
            display_with_box("Q1 Mistakes", ", ".join(q1_mistakes), "red")

        if len(answers) > 1:
            q2_score, q2_similarity, q2_grammar, q2_length, q2_mistakes = evaluate_answer(answers[1], standard_answers['Q2'], 8, length_based=True)
            total_score += q2_score
            all_mistakes["Q2"].extend(q2_mistakes)
            
            display_with_box("Q2 Submitted Score", f"{q2_score:.2f}", "green")
            display_with_box("Q2 Submitted Similarity Marks", f"{q2_similarity:.2f}", "blue")
            display_with_box("Q2 Submitted Grammar Marks", f"{q2_grammar:.2f}", "blue")
            display_with_box("Q2 Submitted Length Marks", f"{q2_length:.2f}", "blue")
            display_with_box("Q2 Mistakes", ", ".join(q2_mistakes), "red")

        display_with_box("Total Submitted Score", f"{total_score:.2f}", "purple")

    if st.session_state.submitted and st.button("Re-evaluate"):
        # Re-evaluation functionality
        answers = st.session_state.answers  # Use stored answers

        total_score = 0
        all_mistakes = {"Q1": [], "Q2": []}

        # Displaying scores, mistakes, and re-evaluated results
        def display_with_box(title, content, color):
            st.markdown(
                f"""
                <div style='border: 1px solid {color}; border-radius: 10px; padding: 10px; margin: 10px 0;'>
                    <strong>{title}</strong>: {content}
                </div>
                """, unsafe_allow_html=True
            )

        if len(answers) > 0:
            q1_score, q1_similarity, q1_grammar, _, q1_mistakes = evaluate_answer(answers[0], standard_answers['Q1'], 2)
            total_score += q1_score
            all_mistakes["Q1"].extend(q1_mistakes)
            
            display_with_box("Q1 Re-evaluated Score", f"{q1_score:.2f}", "green")
            display_with_box("Q1 Re-evaluated Similarity Marks", f"{q1_similarity:.2f}", "blue")
            display_with_box("Q1 Re-evaluated Grammar Marks", f"{q1_grammar:.2f}", "blue")
            display_with_box("Q1 Mistakes", ", ".join(q1_mistakes), "red")

        if len(answers) > 1:
            q2_score, q2_similarity, q2_grammar, q2_length, q2_mistakes = evaluate_answer(answers[1], standard_answers['Q2'], 8, length_based=True)
            total_score += q2_score
            all_mistakes["Q2"].extend(q2_mistakes)
            
            display_with_box("Q2 Re-evaluated Score", f"{q2_score:.2f}", "green")
            display_with_box("Q2 Re-evaluated Similarity Marks", f"{q2_similarity:.2f}", "blue")
            display_with_box("Q2 Re-evaluated Grammar Marks", f"{q2_grammar:.2f}", "blue")
            display_with_box("Q2 Re-evaluated Length Marks", f"{q2_length:.2f}", "blue")
            display_with_box("Q2 Mistakes", ", ".join(q2_mistakes), "red")

        display_with_box("Total Re-evaluated Score", f"{total_score:.2f}", "purple")
